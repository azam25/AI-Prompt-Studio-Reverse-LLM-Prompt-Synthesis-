"""Document processing service for parsing and chunking documents."""
import os
import uuid
from typing import List, Tuple
from pathlib import Path

from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import markdown

from app.config import settings
from app.models.document import Document, DocumentChunk, DocumentType


class DocumentProcessor:
    """Process documents for embedding and retrieval."""
    
    def __init__(self, chunk_size: int = None, chunk_overlap: int = None):
        self.chunk_size = chunk_size or settings.chunk_size
        self.chunk_overlap = chunk_overlap or settings.chunk_overlap
    
    def process_file(self, file_path: str, filename: str) -> Tuple[Document, List[DocumentChunk]]:
        """
        Process a file and return the document with its chunks.
        
        Args:
            file_path: Path to the uploaded file
            filename: Original filename
            
        Returns:
            Tuple of (Document, List[DocumentChunk])
        """
        # Determine document type
        ext = Path(filename).suffix.lower().lstrip('.')
        doc_type = self._get_document_type(ext)
        
        # Extract text based on document type
        text = self._extract_text(file_path, doc_type)
        
        # Create document
        doc_id = str(uuid.uuid4())
        document = Document(
            id=doc_id,
            filename=filename,
            document_type=doc_type,
            file_path=file_path,
            metadata={"original_filename": filename}
        )
        
        # Create chunks
        chunks = self._create_chunks(text, doc_id)
        document.chunk_count = len(chunks)
        
        return document, chunks
    
    def _get_document_type(self, extension: str) -> DocumentType:
        """Map file extension to DocumentType."""
        mapping = {
            'pdf': DocumentType.PDF,
            'docx': DocumentType.DOCX,
            'doc': DocumentType.DOCX,
            'txt': DocumentType.TXT,
            'md': DocumentType.MARKDOWN,
            'markdown': DocumentType.MARKDOWN
        }
        return mapping.get(extension, DocumentType.TXT)
    
    def _extract_text(self, file_path: str, doc_type: DocumentType) -> str:
        """Extract text content from a document."""
        if doc_type == DocumentType.PDF:
            return self._extract_pdf(file_path)
        elif doc_type == DocumentType.DOCX:
            return self._extract_docx(file_path)
        elif doc_type == DocumentType.MARKDOWN:
            return self._extract_markdown(file_path)
        else:
            return self._extract_txt(file_path)
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        return "\n\n".join(text_parts)
    
    def _extract_markdown(self, file_path: str) -> str:
        """Extract text from Markdown file (remove formatting)."""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_text = f.read()
        # Convert to HTML and strip tags for plain text
        # For simplicity, we'll keep the markdown as-is for semantic meaning
        return md_text
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _create_chunks(self, text: str, document_id: str) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks.
        
        Uses a sliding window approach with overlap for context preservation.
        """
        chunks = []
        
        # Clean and normalize text
        text = text.strip()
        if not text:
            return chunks
        
        # Split by sentences for better chunk boundaries
        sentences = self._split_into_sentences(text)
        
        current_chunk = []
        current_length = 0
        chunk_index = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            if current_length + sentence_length > self.chunk_size and current_chunk:
                # Create chunk from current sentences
                chunk_text = " ".join(current_chunk)
                chunk = DocumentChunk(
                    id=f"{document_id}_chunk_{chunk_index}",
                    document_id=document_id,
                    content=chunk_text,
                    chunk_index=chunk_index,
                    metadata={"char_count": len(chunk_text)}
                )
                chunks.append(chunk)
                chunk_index += 1
                
                # Keep overlap sentences
                overlap_text = ""
                overlap_sentences = []
                for s in reversed(current_chunk):
                    if len(overlap_text) + len(s) <= self.chunk_overlap:
                        overlap_sentences.insert(0, s)
                        overlap_text = " ".join(overlap_sentences)
                    else:
                        break
                
                current_chunk = overlap_sentences
                current_length = len(overlap_text)
            
            current_chunk.append(sentence)
            current_length += sentence_length
        
        # Don't forget the last chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunk = DocumentChunk(
                id=f"{document_id}_chunk_{chunk_index}",
                document_id=document_id,
                content=chunk_text,
                chunk_index=chunk_index,
                metadata={"char_count": len(chunk_text)}
            )
            chunks.append(chunk)
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        import re
        # Simple sentence splitting
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]


# Singleton instance
document_processor = DocumentProcessor()


def get_document_processor() -> DocumentProcessor:
    """Get the document processor instance."""
    return document_processor
